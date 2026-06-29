from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "/Users/liuqiyu/Desktop/qveris/06_deliverables/2026-06-29_us_stock_ai_data_requirements_qveris_gap_report_v0.3_explained.docx"
LANDSCAPE_IMAGE = "/Users/liuqiyu/Desktop/qveris/05_data/outputs/visuals/2026-06-29_us_stock_ai_data_capability_landscape.png"
UNIVERSE_IMAGE = "/Users/liuqiyu/Desktop/qveris/05_data/outputs/visuals/2026-06-29_us_stock_ai_data_capability_universe.png"
REPORT_MATRIX_IMAGE = "/Users/liuqiyu/Desktop/qveris/06_deliverables/infographics/financial-data-capability-landscape.png"


SOURCES = [
    {
        "name": "QVeris 首页",
        "url": "https://qveris.com/",
        "use": "确认 QVeris 是面向 AI agent 的 capability routing network，支持 Discover / Inspect / Call，并宣称覆盖 10,000+ capabilities、15+ categories。",
    },
    {
        "name": "QVeris CLI 文档",
        "url": "https://qveris.ai/docs/cli",
        "use": "确认 CLI 可用自然语言 discover 能力、inspect schema/参数/质量/计费、call 执行；示例包含 stock price API、cryptocurrency market data 等。",
    },
    {
        "name": "QVeris Stock Research Agent Guide",
        "url": "https://qveris.ai/guides/build-stock-research-agent-in-cursor/",
        "use": "确认 QVeris 公开描述支持 stock screening、company research、market analysis、market monitoring、earnings and filings research、risk and compliance signals 等金融研究场景。",
    },
    {
        "name": "QVeris Real-Time Stock Price API Guide",
        "url": "https://qveris.ai/guides/real-time-stock-price-api-for-ai-agents/",
        "use": "确认 QVeris 公开讨论 real-time stock price API、quotes、bid-ask spreads、volume、trade history、historical market data、WebSocket trades/quotes、Polygon/Finnhub/Twelve Data/Alpha Vantage 等 provider 路由。",
    },
    {
        "name": "QVeris Pricing",
        "url": "https://qveris.com/pricing",
        "use": "确认 Discover / Inspect 免费，Call 按 credits 消耗；公开示例成本含 live quote、financial report analysis 等。",
    },
    {
        "name": "QVeris Tool Finder/Search",
        "url": "https://qveris.ai/search",
        "use": "确认公开搜索页描述可按工具能力搜索 provider、pricing、latency、success rate；匿名页面未能返回完整金融能力库存，因此细项仍需后续登录 Inspect。",
    },
]


DATA_CATEGORIES = [
    {
        "cat": "1. 市场数据 Market Data",
        "priority": "P0 必备",
        "fields": "实时/延迟股价、历史股价、OHLC、成交量/成交额、盘前盘后、指数/板块表现、行业分类、涨跌幅榜、创新高/新低、52 周高低点、股息/拆股/复权、公司行动。",
        "scenario": "回答“市场现在发生了什么”，为所有个股研究、看板、预警和解释型分析提供时间坐标。",
        "value": "让用户先判断价格变化、成交配合、相对指数/板块强弱，以及变动是否来自市场共振还是个股事件。",
        "features": "个股行情页、watchlist、市场大盘看板、sector rotation、movers 榜单、价格异动解释、基础回测与图表。",
    },
    {
        "cat": "2. 交易与微观结构数据",
        "priority": "P0/P1",
        "fields": "逐笔成交、Bid/Ask、spread、Level 1/Level 2、order book、market depth、order/auction imbalance、trade condition、exchange venue、off-exchange/ATS/TRF。",
        "scenario": "理解价格变化背后的买卖压力、成交质量和交易场所结构。",
        "value": "帮助短线、量化和专业用户识别“有效上涨/下跌”与低质量异动，避免只看收盘价。",
        "features": "交易质量评分、盘口压力、异常成交监控、开收盘竞价观察、微观结构风险提示。",
    },
    {
        "cat": "3. 流动性数据",
        "priority": "P0/P1",
        "fields": "bid-ask spread、effective/quoted spread、ADV、dollar volume、turnover ratio、market depth、Amihud illiquidity、price impact、slippage、block trade、off-exchange ratio。",
        "scenario": "判断股票是否容易交易，以及进出场是否可能造成冲击成本。",
        "value": "对中小盘、低价股、主题股尤其重要，可解释为何账面机会无法低成本执行。",
        "features": "流动性评分、交易可执行性提示、仓位容量估算、风险看板、低流动性预警。",
    },
    {
        "cat": "4. 公司基本面数据",
        "priority": "P0 必备",
        "fields": "收入、毛利/营业利润/净利润、EPS、利润率、EBITDA、经营现金流、自由现金流、CapEx、资产负债表、现金/债务、ROE/ROA/ROIC、R&D、SBC、分业务/地区收入、客户集中度、供应链依赖。",
        "scenario": "判断公司经营质量、盈利能力、增长质量和长期投资价值。",
        "value": "AI/科技股分析不仅看利润，还要看 R&D、CapEx、云业务、数据中心投资和管理层指引。",
        "features": "公司概览页、基本面体检、财务趋势图、盈利质量分析、AI 股核心经营指标卡。",
    },
    {
        "cat": "5. 财报与监管披露数据",
        "priority": "P0 必备",
        "fields": "10-K、10-Q、8-K、S-1、DEF 14A、Form 4、13F、13D/13G、earnings release、investor presentation、press release、risk factors、MD&A、guidance、legal proceedings、buyback disclosure。",
        "scenario": "基于权威披露追溯公司经营、风险、重大事件和管理层表述变化。",
        "value": "建立 AI 产品可信度，减少新闻误读，支撑可溯源的财报/公告摘要。",
        "features": "SEC filing 摘要、风险因素变化追踪、财报问答助手、公告事件抽取、管理层指引对比。",
    },
    {
        "cat": "6. 估值数据",
        "priority": "P0/P1",
        "fields": "Market Cap、EV、P/E、Forward P/E、P/S、EV/Sales、EV/EBITDA、P/B、PEG、FCF Yield、Earnings Yield、Dividend Yield、历史估值区间、估值分位、peer/sector comparison、DCF assumptions、implied growth。",
        "scenario": "回答“股票贵不贵”，并结合增长、盈利质量、行业阶段和利率环境解释估值。",
        "value": "帮助用户避免孤立使用单一倍数，支持 AI 高成长公司的预期和长期 TAM 分析。",
        "features": "估值助手、同行对比、历史分位、隐含增长率、估值与利率敏感性分析。",
    },
    {
        "cat": "7. 分析师预期与市场共识",
        "priority": "P0/P1",
        "fields": "EPS/revenue/EBITDA/FCF consensus、guidance consensus、target price、rating、upgrade/downgrade、estimate revision、earnings/revenue surprise、beat/miss/in-line、consensus trend、dispersion。",
        "scenario": "财报季判断市场原本预期、公司实际交付和管理层指引之间的差异。",
        "value": "解释“业绩好但股价跌”或“业绩一般但股价涨”的预期差逻辑。",
        "features": "预期差分析、财报 surprise 解释、分析师修正监控、rating 变动提醒、财报前后复盘。",
    },
    {
        "cat": "8. 新闻与事件数据",
        "priority": "P0 必备",
        "fields": "公司/行业/宏观新闻、财报日历、产品发布、并购、回购、裁员、诉讼监管、管理层变动、评级调整、大客户订单、合同、政策、地缘政治、供应链新闻、AI infrastructure/semiconductor news。",
        "scenario": "解释短期波动、识别投资催化剂和风险事件。",
        "value": "对 AI 和科技股尤其关键，可连接出口限制、云 capex、模型发布、数据中心订单等事件。",
        "features": "股价异动归因、事件时间线、催化剂雷达、新闻摘要、风险事件提醒。",
    },
    {
        "cat": "9. 宏观经济与利率数据",
        "priority": "P0/P1",
        "fields": "Fed Funds、2Y/10Y Treasury yield、yield curve、CPI/PCE/Core、NFP、unemployment、GDP、PMI、retail sales、consumer confidence、DXY、oil、credit spread、FCI、FOMC calendar/statement/speeches。",
        "scenario": "判断成长股、科技股和高估值股票所处的宏观折现率环境。",
        "value": "帮助解释利率上行/下行对 AI 高估值资产的估值压力或扩张作用。",
        "features": "宏观看板、利率敏感性解释、FOMC 事件提醒、成长股环境评分。",
    },
    {
        "cat": "10. 期权与波动率数据",
        "priority": "P1 增强",
        "fields": "options volume、open interest、IV、historical volatility、IV Rank/Percentile、put/call ratio、unusual options activity、options flow、expected move、gamma/delta exposure、skew、term structure、VIX/VXN/MOVE、earnings implied move。",
        "scenario": "观察市场对未来波动、尾部风险和财报事件的定价。",
        "value": "帮助用户理解财报前 expected move、情绪偏多/偏空和市场风险偏好。",
        "features": "期权情绪面板、财报预期波动、异常期权流、VIX/VXN 风险仪表盘。",
    },
    {
        "cat": "11. 资金流与持仓数据",
        "priority": "P1 增强",
        "fields": "ETF/mutual fund/sector fund flows、institutional ownership、13F holdings、insider transactions、buyback、short interest、days to cover、securities lending、margin debt、hedge fund positioning、retail/smart money flow、fund concentration。",
        "scenario": "观察资金流向、机构偏好、交易拥挤度和潜在反转风险。",
        "value": "对热门 AI 股尤其重要，可识别上涨是否伴随机构拥挤或资金持续流入。",
        "features": "资金流雷达、机构持仓变化、拥挤交易评分、smart money 跟踪、回购与内部人交易监控。",
    },
    {
        "cat": "12. 做空与卖空数据",
        "priority": "P1 增强",
        "fields": "short interest、short interest ratio、days to cover、short volume、short sale volume、borrow fee、shares available to borrow、securities lending utilization、fail to deliver、short squeeze risk。",
        "scenario": "判断负面预期、卖空压力和潜在逼空风险。",
        "value": "对高波动、小盘、meme stock 和主题 AI 股有明显参考价值。",
        "features": "做空风险面板、逼空风险评分、borrow fee 监控、short volume 异动提醒。",
    },
    {
        "cat": "13. 差异化数据：AI 供应链 + Dark Pool/Block/Off-exchange",
        "priority": "P2 差异化",
        "fields": "AI 芯片/GPU/ASIC/AI 加速器出货、HBM、CoWoS、wafer starts、半导体设备订单、AI server、cloud/hyperscaler capex、数据中心建设/电力容量、网络设备/光模块、库存、lead time、backlog、客户/供应商依赖、出口管制；dark pool/ATS/off-exchange volume、block trades、large prints、venue distribution、hidden liquidity、TRF。",
        "scenario": "提前观察 AI 产业真实需求、供给瓶颈、产业链景气度和机构大额交易结构。",
        "value": "形成 QVeris 美股 AI 产品差异化，避免只做通用行情/新闻产品。",
        "features": "AI 供应链追踪、AI 产业链景气指数、云 capex 监控、半导体瓶颈雷达、机构大额交易观察。",
    },
]


COVERAGE = [
    ["市场数据", "实时/历史价格、quotes、bid-ask、volume、trade history、provider 路由", "已公开确认", "QVeris real-time stock price guide 明确讨论 stock price API、quotes、bid-ask、volume、trade history、historical market data；CLI 文档有 stock price API discover 示例。", "需 Inspect 确认具体 provider、覆盖交易所、延迟、盘前盘后、复权口径。", "P0"],
    ["交易与微观结构", "trades、quotes、WebSocket、OPRA/Polygon 提及；Level 2、order book、imbalance 未公开确认", "部分公开确认", "real-time stock price guide 提及 WebSocket trades/quotes 和 Polygon/NASDAQ/OPRA 相关能力。", "Level 2、market depth、auction imbalance、venue 细分需 Inspect；可能需要专门市场数据 provider。", "P0/P1"],
    ["流动性", "spread、volume、trade history 可由行情推导；price impact/slippage/Amihud 未公开确认", "部分公开确认", "公开指南确认 bid-ask spreads、volume、trade history。", "高级流动性指标多数需产品侧计算或另接 provider；block/off-exchange 单列验证。", "P0/P1"],
    ["公司基本面", "fundamentals、company profiles", "已公开确认", "Stock Research Agent guide 将 fundamentals、company profiles 列为碎片化金融数据和研究能力场景。", "需 Inspect 确认字段深度、历史期数、分业务/地区、R&D/SBC/CapEx 等 AI 股关键字段。", "P0"],
    ["财报与监管披露", "filings、earnings and filings research、document extraction", "已公开确认", "Stock Research Agent guide 明确列出 filings、earnings and filings research、company updates and investor materials structured extraction。", "需 Inspect 确认 SEC form 类型、transcript、presentation、press release、Form 4/13F 覆盖。", "P0"],
    ["估值", "market cap/倍数可能由行情+基本面构成；估值分位/peer/DCF 未公开确认", "可能已有但需 Inspect 确认", "公开材料未直接列出 valuation ratios；但行情、基本面可支撑产品侧计算。", "需要确认是否已有 valuation capability；历史估值分位和 peer sets 可能需自建。", "P0/P1"],
    ["分析师预期与共识", "consensus、ratings、estimate revision、target price", "未在公开资料确认", "公开 guide 讨论 market analysis，但未明确 analyst estimates/consensus。", "需重点 Inspect “analyst estimates / consensus / earnings surprise”。若无，需接入专业金融数据源。", "P0/P1"],
    ["新闻与事件", "news、recent news、risk signals、market context", "已公开确认", "Stock Research Agent guide 明确提到 recent news、risk signals、market monitoring；real-time guide 提到 news sentiment/SEC filings 可在同一 reasoning loop 查询。", "需确认新闻源、延迟、事件结构化程度、财报日历、评级调整和供应链新闻覆盖。", "P0"],
    ["宏观经济与利率", "macro market context 可能通过通用金融/数据 provider 获取", "可能已有但需 Inspect 确认", "QVeris 宣称 finance 等多领域 capability，但公开股票 guide 未逐项列出 CPI/Fed/yields。", "需 Inspect “FRED / macro data / treasury yields / FOMC calendar”。", "P0/P1"],
    ["期权与波动率", "OPRA/Polygon 相关提及；VIX/IV/options flow 未公开确认", "部分公开确认", "real-time stock price guide 提到 Polygon 与 NASDAQ/OPRA，说明期权数据入口可能存在。", "需 Inspect options chain、IV、OI、put/call、unusual flow、gamma exposure、VIX/VXN/MOVE。", "P1"],
    ["资金流与持仓", "13F/ownership/insider/buyback/ETF flows", "可能已有但需 Inspect 确认", "公开 guide 提到 risk and compliance signals、company research，但未逐项确认 ownership/fund flow。", "需 Inspect 13F、institutional ownership、ETF flows、insider transactions、buyback 数据。", "P1"],
    ["做空与卖空", "short interest、short volume、borrow fee、FTD", "未在公开资料确认", "公开材料未明确 short interest 或 securities lending。", "需 Inspect FINRA short interest/short sale volume、borrow fee、shares available、FTD。", "P1"],
    ["AI 供应链与 Dark Pool/Block/Off-exchange", "AI supply chain、cloud capex、semicap orders、dark pool/ATS/TRF/block trades", "未在公开资料确认", "公开材料仅泛化提到 finance/risk/market data，未确认这些差异化细项。", "优先作为差异化数据缺口：供应链可能需要行业数据/新闻抽取/专门 provider；dark pool/ATS/TRF 需 FINRA 或专业交易数据源。", "P2"],
]


FEATURES = [
    ["MVP 基础功能", "个股行情与 Watchlist", "市场数据、流动性、新闻事件", "展示实时/延迟价格、涨跌幅、成交、板块归因、新闻解释和基础预警。"],
    ["MVP 基础功能", "公司研究 Agent", "基本面、披露、新闻、估值", "输入 ticker 后输出公司概览、财务趋势、最新披露、关键风险和研究摘要。"],
    ["MVP 基础功能", "财报/披露摘要", "财报披露、分析师预期、新闻", "自动摘要 10-K/10-Q/8-K/earnings release，并标注同比/环比变化、guidance 和风险因素。"],
    ["MVP 基础功能", "市场监控 Dashboard", "市场数据、指数/板块、宏观、新闻", "监控指数、板块、movers、财报日历、FOMC 等关键市场事件。"],
    ["专业增强功能", "预期差分析", "分析师共识、财报实际、价格反应", "解释 beat/miss、estimate revision、target price/rating 变化和股价反应。"],
    ["专业增强功能", "估值助手", "估值、基本面、peer、利率", "做历史分位、同行对比、隐含增长和成长股利率敏感性说明。"],
    ["专业增强功能", "风险 Dashboard", "流动性、期权、做空、资金流、宏观", "汇总高波动、高估值、低流动性、做空拥挤、期权极端情绪等风险。"],
    ["专业增强功能", "机构与资金行为追踪", "13F、ETF flows、insider、buyback、block trades", "识别资金偏好、持仓拥挤、回购支撑和机构大额交易信号。"],
    ["差异化功能", "AI 供应链追踪", "GPU/HBM/CoWoS/cloud capex/data center/semicap", "为 NVDA、AMD、AVGO、MRVL、ASML、AMAT、LRCX 等 AI 链公司建立前瞻景气跟踪。"],
    ["差异化功能", "AI 产业链景气指数", "供应链、新闻、capex、订单、库存、lead time", "把碎片化产业信号合成为需求扩张、供给紧张、库存调整、周期降温等状态。"],
    ["差异化功能", "场外与大额交易观察", "dark pool、ATS、TRF、block trades、venue distribution", "作为资金行为辅助信号，不直接给买卖建议，而用于解释交易结构变化。"],
]


ROADMAP = [
    ["P0 必备", "先验证并接通市场行情、公司基本面、财报披露、新闻事件、基础估值、核心宏观。", "支撑个股研究、watchlist、市场监控、财报摘要四个基础体验。"],
    ["P1 增强", "补齐分析师预期、期权波动率、资金流持仓、做空卖空、高级流动性。", "提升专业度，支撑预期差、风险看板、机构行为和交易拥挤度分析。"],
    ["P2 差异化", "专项建设 AI 供应链和 dark pool/block/off-exchange。", "形成 AI 投资产品辨识度，避免与通用股票工具同质化。"],
]


INSPECT_QUERIES = [
    "real-time US stock quotes with bid ask spread and premarket after-hours",
    "historical OHLCV adjusted prices US equities",
    "company fundamentals income statement balance sheet cash flow US stocks",
    "SEC filings 10-K 10-Q 8-K Form 4 13F",
    "earnings transcripts and earnings releases US public companies",
    "analyst estimates consensus EPS revenue target price rating revisions",
    "US stock news events earnings calendar rating changes",
    "FRED macro data treasury yields CPI PCE FOMC calendar",
    "options chain implied volatility open interest put call ratio unusual options flow",
    "ETF flows institutional ownership insider transactions buybacks",
    "short interest short volume borrow fee fail to deliver",
    "AI supply chain GPU HBM CoWoS cloud capex data center construction",
    "dark pool ATS TRF block trades off exchange volume US equities",
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, name="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_font(r)


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    set_font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True, color="2E74B5" if level < 3 else "1F4D78")
    return p


def add_table(doc, headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_font(r, size=font_size, bold=True)
        set_cell_shading(cell, "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            set_font(r, size=font_size)
    set_table_width(table, widths)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ["List Bullet", "List Number"]:
        st = styles[style_name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(10.5)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.167

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("QVeris US Stock AI Product Data Study | 2026-06-29")
    set_font(r, size=8, color="666666")


def build_doc():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("美股 AI 产品数据需求与 QVeris 能力缺口研究报告")
    set_font(r, size=20, bold=True, color="0B2545")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("版本：v0.3 explained | 日期：2026-06-29 | 范围：公开资料核对，不调用私有账号/API")
    set_font(r, size=10, color="555555")

    add_heading(doc, "摘要", 1)
    add_paragraph(doc, "本报告基于原始文档《美股投资所需数据类型研究报告.docx》，将美股 AI 产品所需数据归纳为 13 个大类，并与 QVeris 官网及公开文档进行能力覆盖核对。")
    add_paragraph(doc, "核心结论：QVeris 公开资料已能确认其适合承载基础金融研究工作流，尤其是市场行情、公司研究、基本面、新闻、财报/披露研究、风险信号和 capability routing；但分析师共识、短空数据、AI 供应链、dark pool/ATS/TRF 等细项尚不能仅凭公开资料确认，需要后续通过 Discover/Inspect 精确验证。")
    add_bullets(doc, [
        "P0 优先建设：行情、公司基本面、财报披露、新闻事件、基础估值、核心宏观数据。",
        "P1 专业增强：分析师预期、期权波动率、资金流持仓、做空卖空、高级流动性。",
        "P2 差异化：AI 供应链和场外/大额交易结构数据，用于形成 AI 投资产品特色。",
        "证据口径：公开资料可确认则标记为“已公开确认/部分公开确认”；无法确认的细项标记为“可能已有但需 Inspect 确认”或“未在公开资料确认”。",
    ])

    add_heading(doc, "可视化总览", 1)
    add_paragraph(doc, "下面三张图用于快速建立全局视角：第一张是报告专属矩阵，展示 13 类美股 AI 数据需求、QVeris 公开覆盖、待 Inspect 缺口与产品工作流之间的关系；后两张是能力全景和分层宇宙图，便于从更宏观的角度理解数据、基础设施和 Agent workflow 的连接。")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("图 1：美股 AI 数据需求与 QVeris 覆盖全景图")
    set_font(r, size=9, bold=True, color="555555")
    doc.add_picture(REPORT_MATRIX_IMAGE, width=Inches(6.5))
    add_heading(doc, "如何阅读图 1 的矩阵", 2)
    add_paragraph(doc, "这张矩阵的左侧是本报告梳理出的 13 类美股 AI 产品数据需求；上方是判断维度，用来回答这类数据是否必要、QVeris 公开资料能否确认、是否可以由产品侧计算、是否还需要 Discover/Inspect，以及它更适合支撑 MVP、专业增强还是差异化功能。")
    add_paragraph(doc, "状态标签含义：")
    add_bullets(doc, [
        "Full：比较明确、强相关，可以作为完整支撑。例如市场数据对 MVP 是 Full，因为个股行情、watchlist 和市场监控都必须依赖它。",
        "Partial：已有一部分公开线索或可间接支持，但字段、覆盖、延迟、供应商、授权或计算口径仍需确认。",
        "Missing：公开资料未确认。这里不等于 QVeris 一定没有，只表示不能仅凭官网和公开文档证明已有。",
        "N/A：不适用于该阶段或该功能层。例如部分专业数据并不是 MVP 第一版必须项。",
    ])
    add_paragraph(doc, "列含义：")
    add_bullets(doc, [
        "数据需求 Required Need：这类数据对美股 AI 产品是否必要。13 类数据基本都是 Full，因为它们共同构成完整投资研究数据栈。",
        "QVeris公开 Public Coverage：从 QVeris 官网、公开文档和 guides 是否能确认相关能力。",
        "产品可计算 Product Computed：即使没有直接字段，产品侧是否能由行情、基本面、披露等数据计算出指标，例如估值倍数或流动性评分。",
        "待Inspect Inspect Needed：后续是否需要用 QVeris Discover/Inspect 验证 tool_id、provider、schema、latency、success rate 和 billing rule。",
        "MVP支撑 MVP Support：是否直接支撑第一版基础产品，例如个股行情、公司研究、财报摘要、新闻解释和市场监控。",
        "专业增强 Advanced：是否更适合专业功能，例如预期差、期权波动、资金流、做空、机构行为和高级风险分析。",
        "差异化 Differentiation：是否能形成 QVeris 美股 AI 产品特色，尤其是 AI 供应链、云 CapEx、半导体瓶颈、dark pool / block trade / off-exchange 信号。",
    ])
    add_paragraph(doc, "阅读示例：市场数据这一行在“QVeris公开”和“MVP支撑”上是 Full，说明公开资料已经能确认 QVeris 具备股票行情相关能力，而且它是第一版产品的底座；但在“产品可计算”和“待Inspect”上仍为 Partial，因为盘前盘后、复权口径、供应商覆盖、延迟和字段细节仍需进一步确认。AI 供应链/场外这一行在“QVeris公开”上是 Missing、在“差异化”上是 Full，说明它是产品特色方向，但公开资料尚未证明 QVeris 已覆盖，需要后续重点 Inspect 或补充数据源。")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("图 2：美股 AI 数据能力全景图")
    set_font(r, size=9, bold=True, color="555555")
    doc.add_picture(LANDSCAPE_IMAGE, width=Inches(6.5))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("图 3：美股 AI 数据能力宇宙图")
    set_font(r, size=9, bold=True, color="555555")
    doc.add_picture(UNIVERSE_IMAGE, width=Inches(6.5))

    add_heading(doc, "一、研究目标与证据口径", 1)
    add_paragraph(doc, "研究目标不是直接定义最终产品，而是回答：如果 QVeris 要做面向美股和 AI 投资场景的产品，需要哪些数据；QVeris 公开资料已能证明哪些能力；缺口应如何转化为后续产品与数据建设任务。")
    add_paragraph(doc, "本报告不调用私有 QVeris 账号、API key 或付费 capability。所有 QVeris 覆盖判断仅基于官网、公开文档、公开 guides、pricing 和公开 Tool Finder 页面。")
    add_paragraph(doc, "覆盖状态定义：")
    add_bullets(doc, [
        "已公开确认：官方公开页面明确提到该类数据或直接相关金融能力。",
        "部分公开确认：官方公开页面确认基础数据或邻近能力，但高级字段/完整覆盖需进一步确认。",
        "可能已有但需 Inspect 确认：QVeris 平台形态或相邻能力支持该方向，但公开资料未列出明确字段。",
        "未在公开资料确认：公开资料没有足够证据证明该数据已经可用。",
    ])

    add_heading(doc, "二、美股 AI 产品所需数据分类", 1)
    for item in DATA_CATEGORIES:
        add_heading(doc, item["cat"], 2)
        add_paragraph(doc, f"优先级：{item['priority']}")
        add_paragraph(doc, f"核心字段：{item['fields']}")
        add_paragraph(doc, f"使用场景：{item['scenario']}")
        add_paragraph(doc, f"用户价值：{item['value']}")
        add_paragraph(doc, f"产品功能映射：{item['features']}")

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "三、QVeris 公开能力覆盖矩阵", 1)
    add_paragraph(doc, "下表按公开资料核对 QVeris 当前能力覆盖。需要注意：QVeris 是 capability routing network，公开页面通常描述能力类别与工作流，并不等价于完整 provider inventory；因此未公开确认不代表实际没有，只代表需要后续通过 Discover/Inspect 验证。")
    add_table(
        doc,
        ["数据类别", "公开可见能力/线索", "覆盖判断", "公开证据", "主要缺口/后续验证", "优先级"],
        COVERAGE,
        [0.85, 1.25, 0.85, 1.55, 1.65, 0.35],
        font_size=7.4,
    )

    add_heading(doc, "四、数据到产品功能的映射", 1)
    add_paragraph(doc, "产品功能应从“数据能回答什么问题”出发，而不是简单堆字段。基础功能保证用户能完成日常研究，增强功能提升专业深度，差异化功能决定 QVeris 美股 AI 产品是否有独特定位。")
    add_table(
        doc,
        ["层级", "功能", "依赖数据", "功能说明"],
        FEATURES,
        [1.0, 1.25, 1.7, 2.55],
        font_size=8,
    )

    add_heading(doc, "五、优先级路线图", 1)
    add_table(
        doc,
        ["阶段", "数据建设重点", "支撑产品目标"],
        ROADMAP,
        [0.95, 2.75, 2.8],
        font_size=8.5,
    )
    add_paragraph(doc, "建议执行顺序：先完成 P0 数据能力的 QVeris Discover/Inspect 盘点，明确 provider、字段、价格、延迟、覆盖范围和返回结构；再设计 MVP 功能；随后补齐 P1 专业增强数据；最后用 P2 差异化数据建立 AI 投资专属壁垒。")

    add_heading(doc, "六、待 Inspect 查询清单", 1)
    add_paragraph(doc, "后续拿到 QVeris 账号或 API key 后，建议按以下查询逐项 Discover/Inspect，并记录 tool_id、provider、schema、latency、success rate、billing rule、region/coverage、字段样例。")
    add_numbered(doc, INSPECT_QUERIES)

    add_heading(doc, "七、后续数据补齐建议", 1)
    add_bullets(doc, [
        "对 P0 数据，优先确认是否已有稳定、低延迟、覆盖美股主流股票的 provider；若多 provider 可用，应比较字段完整性、成本和成功率。",
        "对分析师共识、评级、target price、earnings surprise，若 QVeris 暂无明确能力，需要评估专业金融数据 provider。",
        "对宏观利率，可优先确认 FRED、Treasury、BLS、BEA、FOMC 日历与声明类能力。",
        "对期权与波动率，需区分基础 options chain、OPRA 行情、IV/OI、unusual flow、gamma exposure，这些数据成本和授权差异较大。",
        "对 AI 供应链，短期可用新闻/披露抽取和行业报告摘要搭建半结构化能力；中期再接入专门供应链、半导体、数据中心和云 capex provider。",
        "对 dark pool/ATS/TRF，必须标明延迟、解释限制和非买卖信号属性，避免产品误导用户。",
    ])

    add_heading(doc, "附录 A：QVeris 官方公开来源", 1)
    add_table(
        doc,
        ["来源", "链接", "本报告用途"],
        [[s["name"], s["url"], s["use"]] for s in SOURCES],
        [1.25, 2.15, 3.1],
        font_size=8,
    )

    add_heading(doc, "附录 B：结论摘要", 1)
    add_paragraph(doc, "QVeris 美股 AI 产品的数据建设应围绕“基础完整性、专业可信度、AI 方向差异化”三条线推进。基础层要让用户完成常见投资研究问题；增强层要解释预期差、波动、资金和风险；差异化层要抓住 AI 供应链和机构交易结构这两类更能体现产品特色的数据。")
    add_paragraph(doc, "本报告的最大限制是 QVeris 公开页面没有暴露完整 capability inventory，因此它是产品和数据规划 v0.1，不是最终采购或工程接入清单。下一步需要用 Discover/Inspect 生成可执行的 tool-level 数据矩阵。")

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
